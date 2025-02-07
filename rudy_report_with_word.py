import streamlit as st
import xml.etree.ElementTree as ET
import re  # For regex operations in parse_xml
from collections import defaultdict
import pandas as pd
import math  # For grid dimensions in the country table

# ---------------------------
# XML Parsing Function
# ---------------------------
def parse_xml(uploaded_file):
    try:
        content = uploaded_file.read()
        content_str = content.decode("utf-8-sig").strip()
        # Remove non-printable characters (except tabs and newlines)
        content_str = re.sub(r"[^\x20-\x7E\t\n\r]", "", content_str)
        root = ET.fromstring(content_str)
        return root
    except UnicodeDecodeError:
        st.error("Encoding issue: Ensure the file is saved as UTF-8.")
    except ET.ParseError as e:
        st.error(f"XML Parsing Error: {e}")
        st.write("Problematic XML Content (after cleanup):", content_str[:500])
    return None

# ---------------------------
# Data Extraction Functions
# ---------------------------
def build_fencer_dict_and_team_dict(root):
    fencer_dict = {}
    team_dict = {}
    for equipe in root.findall(".//Equipe"):
        team_id = equipe.get('ID')
        nation = equipe.get('Nation')
        team_name = equipe.get('IdOrigine') if team_id and team_id.isdigit() else team_id
        if not team_id or not nation or not team_name:
            continue

        team_dict[team_id] = {"Team Name": team_name, "Nation": nation}

        for tireur in equipe.findall(".//Tireur"):
            fencer_id = tireur.get('ID')
            fencer_name = f"{tireur.get('Prenom')} {tireur.get('Nom')}"
            date_of_birth = tireur.get("DateNaissance", "Unknown")
            lateralite = tireur.get("Lateralite", "Unknown")

            fencer_dict[fencer_id] = {
                "Name": fencer_name,
                "Date of Birth": date_of_birth,
                "Lateralite": lateralite,
                "EquipeID": team_id,
                "Nation": nation,
                "Team Name": team_name
            }
    return fencer_dict, team_dict

def extract_final_rankings(root, team_dict):
    rankings = []
    for phase in root.findall(".//PhaseDeTableaux"):
        for equipe in phase.findall(".//Equipe"):
            team_id = equipe.get('REF')
            final_rank = equipe.get('RangFinal')
            if team_id and final_rank:
                team_name = team_dict.get(team_id, {}).get("Team Name", team_id)
                nation = team_dict.get(team_id, {}).get("Nation", "Unknown")
                rankings.append({
                    "Team Name": team_name,
                    "Nation": nation,
                    "Final Rank": int(final_rank)
                })
    return pd.DataFrame(rankings).sort_values(by="Final Rank").reset_index(drop=True)

def generate_tables_data(root, team_dict, fencer_dict):
    results_by_stage = defaultdict(list)
    accumulated_touches = defaultdict(lambda: {"scored": 0, "against": 0, "matches": []})

    for suite in root.findall(".//SuiteDeTableaux"):
        for tableau in suite.findall(".//Tableau"):
            stage_title = tableau.get('Titre', 'Unknown Stage')
            for match in tableau.findall(".//Match"):
                team_d_ref = match.find(".//Equipe[@Cote='D']")
                team_g_ref = match.find(".//Equipe[@Cote='G']")
                team_d = team_dict.get(team_d_ref.get('REF')) if team_d_ref is not None else None
                team_g = team_dict.get(team_g_ref.get('REF')) if team_g_ref is not None else None
                team_d_name = team_d["Team Name"] if team_d else "Unknown Team"
                team_g_name = team_g["Team Name"] if team_g else "Unknown Team"

                # Only process matches involving Qatar ("QAT")
                if (team_d and team_d["Nation"] == "QAT") or (team_g and team_g["Nation"] == "QAT"):
                    prev_score_1, prev_score_2 = 0, 0
                    for assaut in match.findall(".//Assaut"):
                        fencer_d_ref = assaut.find(".//Tireur[@Cote='D']").get('REF')
                        fencer_g_ref = assaut.find(".//Tireur[@Cote='G']").get('REF')
                        score_d = int(assaut.find(".//Tireur[@Cote='D']").get('Score', 0))
                        score_g = int(assaut.find(".//Tireur[@Cote='G']").get('Score', 0))
                        touches_d = score_d - prev_score_1
                        touches_g = score_g - prev_score_2
                        diff_in_touches = touches_d - touches_g
                        prev_score_1, prev_score_2 = score_d, score_g

                        fencer_d = fencer_dict.get(fencer_d_ref, {"Name": "Unknown", "Nation": "Unknown"})
                        fencer_g = fencer_dict.get(fencer_g_ref, {"Name": "Unknown", "Nation": "Unknown"})

                        if fencer_d["Nation"] == "QAT":
                            accumulated_touches[fencer_d["Name"]]["scored"] += touches_d
                            accumulated_touches[fencer_d["Name"]]["against"] += touches_g
                            outcome = "Victory" if touches_d > touches_g else "Defeat" if touches_d < touches_g else "Draw"
                            accumulated_touches[fencer_d["Name"]]["matches"].append({
                                "Opponent Team": team_g_name,
                                "Outcome": outcome
                            })

                        if fencer_g["Nation"] == "QAT":
                            accumulated_touches[fencer_g["Name"]]["scored"] += touches_g
                            accumulated_touches[fencer_g["Name"]]["against"] += touches_d
                            outcome = "Victory" if touches_g > touches_d else "Defeat" if touches_g < touches_d else "Draw"
                            accumulated_touches[fencer_g["Name"]]["matches"].append({
                                "Opponent Team": team_d_name,
                                "Outcome": outcome
                            })

                        row_data = {
                            "Team_1": team_d_name,
                            "Fencer_1": fencer_d["Name"],
                            "Touches_1": f"{touches_d} ({diff_in_touches})",
                            "Score_1": score_d,
                            "Score_2": score_g,
                            "Touches_2": f"{touches_g} ({-diff_in_touches})",
                            "Fencer_2": fencer_g["Name"],
                            "Team_2": team_g_name
                        }
                        results_by_stage[stage_title].append(row_data)
    return results_by_stage, accumulated_touches

def generate_qatari_summary(accumulated_touches):
    qatari_summary = []
    for fencer, scores in accumulated_touches.items():
        total = scores["scored"] - scores["against"]
        qatari_summary.append({
            "Fencer": fencer,
            "Scored": scores["scored"],
            "Conceded": scores["against"],
            "Total": total,
            "Matches": scores["matches"]
        })
    return qatari_summary

def get_qatari_fencers_table(fencer_dict):
    # Get all fencers whose Nation is "QAT"
    qatari_fencers = [fencer for fencer in fencer_dict.values() if fencer.get("Nation") == "QAT"]
    df = pd.DataFrame(qatari_fencers)
    # Ensure the expected columns exist even if no data is present.
    for col in ["Name", "Date of Birth"]:
        if col not in df.columns:
            df[col] = pd.Series(dtype='str')
    return df

def get_country_team_counts(team_dict):
    country_team_count = defaultdict(int)
    for details in team_dict.values():
        nation = details['Nation']
        if nation and len(nation) == 3:
            country_team_count[nation] += 1
    team_count_df = pd.DataFrame(list(country_team_count.items()), columns=["Country", "Number of Teams"])
    return len(team_dict), len({d['Nation'] for d in team_dict.values()}), team_count_df

# ---------------------------
# Streamlit App Main Logic
# ---------------------------
st.sidebar.header("Upload XML File")
uploaded_file = st.sidebar.file_uploader("Choose an XML file", type="xml")

if uploaded_file:
    root = parse_xml(uploaded_file)
    if root is None:
        st.stop()  # Stop if XML cannot be parsed

    # Extract overview information from XML
    year = root.get("Annee", "Unknown Year")
    tournament = root.get("TitreCourtTournoi") or root.get("TitreCourt") or "Unknown Tournament"
    championship = root.get("Championnat", "Unknown Championship")
    category = "Cadet" if root.get("Categorie") == "C" else "Junior" if root.get("Categorie") == "J" else root.get("Categorie")
    weapon = "Epee" if root.get("Arme") == "E" else "Sabre" if root.get("Arme") == "S" else "Foil" if root.get("Arme") == "F" else root.get("Arme")
    gender = "Male" if root.get("Sexe") == "M" else "Female" if root.get("Sexe") == "F" else root.get("Sexe")
    date = root.get("Date", "Unknown Date")
    location = root.get("Lieu", "Unknown Location")

    # Build dictionaries and DataFrames
    fencer_dict, team_dict = build_fencer_dict_and_team_dict(root)
    final_rankings_df = extract_final_rankings(root, team_dict)
    tables_data, accumulated_touches = generate_tables_data(root, team_dict, fencer_dict)
    qatari_summary = generate_qatari_summary(accumulated_touches)
    qatari_fencers_df = get_qatari_fencers_table(fencer_dict)[["Name", "Date of Birth"]]
    num_teams, num_countries, team_count_df = get_country_team_counts(team_dict)

    # ---------------------------
    # Display App Tabs
    # ---------------------------
    tabs = st.tabs(["Overview", "Nation Overview", "Tables", "Review", "Final Rankings"])

    with tabs[0]:
        st.header(f"{year} - {tournament} - {championship}")
        st.subheader(f"Date: {date}, Location: {location}")
        st.subheader(f"Category: {category}, Weapon: {weapon}, Gender: {gender}")
        st.markdown("---")
        qatar_ranking = final_rankings_df[final_rankings_df['Nation'] == "QAT"]
        if not qatar_ranking.empty:
            qatar_final_rank = qatar_ranking.iloc[0]['Final Rank']
            st.subheader(f"The Qatar team achieved a final ranking of {qatar_final_rank}.")
        else:
            st.subheader("No final ranking data available for the Qatar team.")
        if not qatari_fencers_df.empty:
            st.markdown("### Qatari Fencers who participated")
            st.table(qatari_fencers_df)
        else:
            st.write("No Qatari fencers found in this competition.")
        st.markdown("---")
        st.subheader(f"**Number of Teams:** {num_teams}")
        st.subheader(f"**Number of Countries:** {num_countries}")
        st.markdown("### Number of Teams per Country")
        num_columns = 4
        rows = (len(team_count_df) + num_columns - 1) // num_columns
        for row_idx in range(rows):
            cols = st.columns(num_columns)
            for col_idx, col in enumerate(cols):
                if row_idx + col_idx * rows < len(team_count_df):
                    country, count = team_count_df.iloc[row_idx + col_idx * rows]
                    col.metric(country, count)

    with tabs[1]:
        nations = sorted({details['Nation'] for details in team_dict.values()})
        selected_nation = st.selectbox("Select a Nation", nations)
        if selected_nation:
            fencer_data = [fencer for fencer in fencer_dict.values() if fencer['Nation'] == selected_nation]
            if fencer_data:
                st.table(pd.DataFrame(fencer_data))
            else:
                st.write(f"No fencers found for {selected_nation}.")

    with tabs[2]:
        for stage_title, matches in tables_data.items():
            st.subheader(f"Stage: {stage_title}")
            st.table(matches)

    with tabs[3]:
        if qatari_summary:
            for fencer_data in qatari_summary:
                st.markdown(f"### Fencer: {fencer_data['Fencer']}")
                st.write(f"Scored: {fencer_data['Scored']}, Conceded: {fencer_data['Conceded']}, Total: {fencer_data['Total']}")
                match_outcomes_df = pd.DataFrame(fencer_data['Matches'])
                if not match_outcomes_df.empty:
                    outcome_summary = (
                        match_outcomes_df
                        .groupby(['Opponent Team', 'Outcome'])
                        .size()
                        .unstack(fill_value=0)
                        .reindex(columns=["Victory", "Defeat", "Draw"], fill_value=0)
                        .reset_index()
                    )
                    st.table(outcome_summary)
                else:
                    st.write("No match outcome data available for this fencer.")
        else:
            st.write("No match outcome data available for Qatari fencers.")

    with tabs[4]:
        st.header("Final Rankings")
        if not final_rankings_df.empty:
            def highlight_qatar(row):
                return ['background-color: yellow' if row['Nation'] == "QAT" else '' for _ in row]
            st.dataframe(final_rankings_df.style.apply(highlight_qatar, axis=1))
        else:
            st.write("No final rankings data available.")

    # ---------------------------
    # Word Export Functions with Styling
    # ---------------------------
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO

    # Define add_heading so it can be used below
    def add_heading(doc, text, level=1):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.size = Pt(16) if level == 1 else Pt(14)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set paragraph text color to black
    def set_paragraph_text_black(paragraph):
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Add a table with header and striped rows styling
    def add_table(doc, df, striped=False):
        HEADER_MAROON = '8A1538'
        LIGHT_MAROON = 'AD5B74'
        WHITE = 'FFFFFF'
        
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Header row styling
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            tcPr = hdr_cells[i]._element.get_or_add_tcPr()
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), HEADER_MAROON)
            tcPr.append(shading_elm)
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Apply striped (banded row) styling if requested
        if striped:
            for i, row in enumerate(table.rows):
                if i == 0:
                    continue
                fill_color = LIGHT_MAROON if (i % 2 == 0) else WHITE
                for cell in row.cells:
                    cell_props = cell._element.get_or_add_tcPr()
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), fill_color)
                    cell_props.append(shading)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
        return table

    # Format cells in an existing table (if needed)
    def format_table_cells(table, font_size=8, background_color="D9E1F2"):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)
                cell_properties = cell._element.find(qn("w:tcPr"))
                if cell_properties is None:
                    cell_properties = OxmlElement("w:tcPr")
                    cell._element.append(cell_properties)
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:val"), "clear")
                shading_elm.set(qn("w:color"), "auto")
                shading_elm.set(qn("w:fill"), background_color)
                cell_properties.append(shading_elm)

    # Create a grid table for "Number of Teams per Country"
    def add_country_team_table(doc, df, columns_per_row=4):
        num_rows = (len(df) + columns_per_row - 1) // columns_per_row
        table = doc.add_table(rows=num_rows, cols=columns_per_row * 2, style="Table Grid")
        for idx, (country, count) in enumerate(df.itertuples(index=False)):
            row_idx = idx // columns_per_row
            col_idx = (idx % columns_per_row) * 2
            table.cell(row_idx, col_idx).text = country
            table.cell(row_idx, col_idx + 1).text = str(count)
            format_table_cells(table, font_size=8)
        doc.add_paragraph()  # Space after the table

    # ---------------------------
    # Export to Word Document Function
    # ---------------------------
    def export_to_word(overview_data, qatari_fencers_df, team_count_df, tables_data, qatari_summary, final_rankings_df):
        doc = Document()
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Competition Overview Section
        add_heading(doc, "Competition Overview", level=1)
        overview_text = (
            f"{overview_data['Year']} - {overview_data['Tournament']} - {overview_data['Championship']}\n"
            f"Date: {overview_data['Date']}, Location: {overview_data['Location']}\n"
            f"Category: {overview_data['Category']}, Weapon: {overview_data['Weapon']}, Gender: {overview_data['Gender']}\n"
            f"Number of Teams: {overview_data['Num Teams']}, Number of Countries: {overview_data['Num Countries']}"
        )
        para = doc.add_paragraph(overview_text)
        set_paragraph_text_black(para)

        # Qatari Fencers Participating Section
        if not qatari_fencers_df.empty:
            add_heading(doc, "Qatari Fencers Participating", level=2)
            add_table(doc, qatari_fencers_df, striped=True)
        
        # Number of Teams per Country Section
        add_heading(doc, "Number of Teams per Country", level=2)
        add_country_team_table(doc, team_count_df)

        # Page break before Tables Overview
        doc.add_page_break()
        add_heading(doc, "Tables Overview", level=1)
        table_count = 0
        for stage, table_data in tables_data.items():
            if table_count % 2 == 0 and table_count > 0:
                doc.add_page_break()
            add_table(doc, pd.DataFrame(table_data), striped=True)
            table_count += 1

        # Page break before Review Section
        doc.add_page_break()
        add_heading(doc, "Review - Accumulated Touches and Match Outcomes for Qatari Fencers", level=1)
        for fencer_data in qatari_summary:
            fencer_text = (
                f"Fencer: {fencer_data['Fencer']}\n"
                f"Scored: {fencer_data['Scored']}, Conceded: {fencer_data['Conceded']}, Total: {fencer_data['Total']}\n"
            )
            para = doc.add_paragraph(fencer_text, style="Normal")
            set_paragraph_text_black(para)
            match_outcomes_df = pd.DataFrame(fencer_data['Matches'])
            if not match_outcomes_df.empty:
                outcome_summary = (
                    match_outcomes_df
                    .groupby(['Opponent Team', 'Outcome'])
                    .size()
                    .unstack(fill_value=0)
                    .reindex(columns=["Victory", "Defeat", "Draw"], fill_value=0)
                    .reset_index()
                )
                add_table(doc, outcome_summary, striped=True)

        # Page break before Final Rankings
        doc.add_page_break()
        add_heading(doc, "Final Rankings", level=1)
        add_table(doc, final_rankings_df, striped=True)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # ---------------------------
    # Word Document Download Button
    # ---------------------------
    with st.sidebar:
        st.header("Download Report")
        overview_data = {
            "Year": year,
            "Tournament": tournament,
            "Championship": championship,
            "Category": category,
            "Weapon": weapon,
            "Gender": gender,
            "Date": date,
            "Location": location,
            "Num Teams": num_teams,
            "Num Countries": num_countries
        }
        word_buffer = export_to_word(
            overview_data=overview_data,
            qatari_fencers_df=qatari_fencers_df,
            team_count_df=team_count_df,
            tables_data=tables_data,
            qatari_summary=qatari_summary,
            final_rankings_df=final_rankings_df
        )
        
        st.download_button(
            label="Download Word Document",
            data=word_buffer,
            file_name="Competition_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("Please upload an XML file to see the report and download the document.")