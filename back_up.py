import streamlit as st
import xml.etree.ElementTree as ET
from collections import defaultdict
import pandas as pd
from io import BytesIO
from docx import Document

# Set the page configuration
st.set_page_config(layout="wide")

# Function to extract team details and build a fencer dictionary
def build_fencer_dict_and_team_dict(root):
    fencer_dict = {}
    team_dict = {}
    for equipe in root.findall(".//Equipe"):
        team_id = equipe.get('ID')
        nation = equipe.get('Nation')
        
        if team_id and team_id.isdigit():
            team_name = equipe.get('IdOrigine')
        else:
            team_name = team_id

        if not team_id or not nation or not team_name:
            continue

        team_dict[team_id] = {
            "Team Name": team_name,
            "Nation": nation
        }

        for tireur in equipe.findall(".//Tireur"):
            fencer_id = tireur.get('ID')
            fencer_name = f"{tireur.get('Prenom')} {tireur.get('Nom')}"
            date_of_birth = tireur.get("DateNaissance", "Unknown")
            lateralite = tireur.get("Lateralite", "Unknown")

            fencer_dict[fencer_id] = {
                "Name": fencer_name,
                "Date of Birth": date_of_birth,
                "Lateralite": lateralite,
                "EquipeID": team_id,
                "Nation": nation,
                "Team Name": team_name
            }
    return fencer_dict, team_dict

# Function to parse rankings and create DataFrame
def extract_final_rankings(root, team_dict):
    rankings = []
    for phase in root.findall(".//PhaseDeTableaux"):
        for equipe in phase.findall(".//Equipe"):
            team_id = equipe.get('REF')
            final_rank = equipe.get('RangFinal')
            
            if team_id and final_rank:
                team_name = team_dict.get(team_id, {}).get("Team Name", team_id)
                nation = team_dict.get(team_id, {}).get("Nation", "Unknown")

                rankings.append({
                    "Team Name": team_name,
                    "Nation": nation,
                    "Final Rank": int(final_rank)
                })

    rankings_df = pd.DataFrame(rankings).sort_values(by="Final Rank").reset_index(drop=True)
    return rankings_df

# Function to create and download Word document
def create_word_document(overview_data, qatari_fencers, tables_data, review_data, final_rankings_data):
    doc = Document()
    
    # Page 1: Overview Data
    doc.add_heading("Competition Overview", level=1)
    for key, value in overview_data.items():
        doc.add_paragraph(f"{key}: {value}")
    
    doc.add_paragraph("Qatari Fencers Participating:", style="Heading 2")
    if qatari_fencers:
        qatari_df = pd.DataFrame(qatari_fencers)
        table = doc.add_table(rows=1, cols=len(qatari_df.columns))
        for i, col_name in enumerate(qatari_df.columns):
            table.rows[0].cells[i].text = col_name
        for _, row in qatari_df.iterrows():
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)
    else:
        doc.add_paragraph("No Qatari fencers found in this competition.")
    
    doc.add_page_break()

    # Pages 2-4: Tables Data (2 tables per page)
    doc.add_heading("Tables Overview", level=1)
    tables_per_page = 2
    for idx, (stage_title, matches) in enumerate(tables_data.items()):
        doc.add_heading(f"Stage: {stage_title}", level=2)
        if isinstance(matches, list) and matches:
            table_df = pd.DataFrame(matches)
            table = doc.add_table(rows=1, cols=len(table_df.columns))
            for i, col_name in enumerate(table_df.columns):
                table.rows[0].cells[i].text = col_name
            for _, row in table_df.iterrows():
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    row_cells[i].text = str(cell)
            if (idx + 1) % tables_per_page == 0:
                doc.add_page_break()

    # Page 5: Review Data
    doc.add_heading("Review - Accumulated Touches and Match Outcomes for Qatari Fencers", level=1)
    for fencer, fencer_data in review_data.items():
        doc.add_paragraph(f"Fencer: {fencer}", style='Heading 2')
        doc.add_paragraph(f"Scored: {fencer_data['Scored']}, Conceded: {fencer_data['Conceded']}, Total: {fencer_data['Total']}")
        
        if fencer_data['Matches']:
            match_outcomes_df = pd.DataFrame(fencer_data['Matches'])
            outcome_summary = (
                match_outcomes_df
                .groupby(['Opponent Team', 'Outcome'])
                .size()
                .unstack(fill_value=0)
                .reset_index()
            )
            table = doc.add_table(rows=1, cols=len(outcome_summary.columns))
            for i, col_name in enumerate(outcome_summary.columns):
                table.rows[0].cells[i].text = col_name
            for _, row in outcome_summary.iterrows():
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    row_cells[i].text = str(cell)
    doc.add_page_break()

    # Last Page: Final Rankings
    doc.add_heading("Final Rankings", level=1)
    table = doc.add_table(rows=1, cols=len(final_rankings_data.columns))
    for i, col_name in enumerate(final_rankings_data.columns):
        table.rows[0].cells[i].text = col_name
    for _, row in final_rankings_data.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)

    # Save to a buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Sidebar file uploader
st.sidebar.header("Upload XML File")
uploaded_file = st.sidebar.file_uploader("Choose an XML file", type="xml")

if uploaded_file:
    tree = ET.parse(uploaded_file)
    root = tree.getroot()
    
    # Overview Tab Data
    overview_data = {
        "Year": root.get("Annee", "Unknown Year"),
        "Tournament": root.get("TitreCourtTournoi") or root.get("TitreCourt") or "Unknown Tournament",
        "Championship": root.get("Championnat", "Unknown Championship"),
        "Category": root.get("Categorie", "Unknown Category"),
        "Weapon": root.get("Arme", "Unknown Weapon"),
        "Gender": root.get("Sexe", "Unknown Gender"),
        "Date": root.get("Date", "Unknown Date"),
        "Location": root.get("Lieu", "Unknown Location")
    }

    # Build the fencer and team dictionaries
    fencer_dict, team_dict = build_fencer_dict_and_team_dict(root)

    # Final Rankings Data
    final_rankings_df = extract_final_rankings(root, team_dict)

    # Get Qatari fencers
    qatari_fencers = [fencer for fencer in fencer_dict.values() if fencer["Nation"] == "QAT"]

    # Tabs for different sections
    tabs = st.tabs(["Overview", "Nation Overview", "Tables", "Review", "Final Rankings"])

    with tabs[0]:
        # Overview Tab
        st.header("Competition Overview")
        for key, value in overview_data.items():
            st.write(f"{key}: {value}")
        
        if qatari_fencers:
            st.write("Qatari Fencers Participating:")
            st.table(pd.DataFrame(qatari_fencers))
        else:
            st.write("No Qatari fencers found in this competition.")

    with tabs[1]:
        # Nation Overview Tab
        st.header("Nation Overview")
        nations = sorted({details['Nation'] for details in team_dict.values()})
        selected_nation = st.selectbox("Select a Nation", nations)

        if selected_nation:
            fencer_data = [fencer for fencer in fencer_dict.values() if fencer['Nation'] == selected_nation]
            if fencer_data:
                st.table(pd.DataFrame(fencer_data))
            else:
                st.write(f"No fencers found for {selected_nation}.")

    tables_data = defaultdict(list)
    with tabs[2]:
        # Tables Tab
        st.header("Tables Overview")
        for suite in root.findall(".//SuiteDeTableaux"):
            for tableau in suite.findall(".//Tableau"):
                stage_title = tableau.get('Titre', 'Unknown Stage')
                stage_data = []
                for match in tableau.findall(".//Match"):
                    team_d_ref = match.find(".//Equipe[@Cote='D']")
                    team_g_ref = match.find(".//Equipe[@Cote='G']")

                    team_d = team_dict.get(team_d_ref.get('REF')) if team_d_ref is not None else None
                    team_g = team_dict.get(team_g_ref.get('REF')) if team_g_ref is not None else None

                    team_d_name = team_d["Team Name"] if team_d else "Unknown Team"
                    team_g_name = team_g["Team Name"] if team_g else "Unknown Team"

                    row_data = {
                        "Team 1": team_d_name,
                        "Score 1": match.find(".//Tireur[@Cote='D']").get('Score', 0),
                        "Score 2": match.find(".//Tireur[@Cote='G']").get('Score', 0),
                        "Team 2": team_g_name
                    }
                    stage_data.append(row_data)
                tables_data[stage_title].append(stage_data)
                st.subheader(f"Stage: {stage_title}")
                st.table(pd.DataFrame(stage_data))

    review_data = defaultdict(lambda: {"Scored": 0, "Conceded": 0, "Total": 0, "Matches": []})
    with tabs[3]:
        # Review Tab
        st.header("Review - Accumulated Touches and Match Outcomes for Qatari Fencers")
        for suite in root.findall(".//SuiteDeTableaux"):
            for tableau in suite.findall(".//Tableau"):
                for match in tableau.findall(".//Match"):
                    for assaut in match.findall(".//Assaut"):
                        fencer_d_ref = assaut.find(".//Tireur[@Cote='D']").get('REF')
                        fencer_g_ref = assaut.find(".//Tireur[@Cote='G']").get('REF')
                        score_d = int(assaut.find(".//Tireur[@Cote='D']").get('Score', 0))
                        score_g = int(assaut.find(".//Tireur[@Cote='G']").get('Score', 0))

                        fencer_d = fencer_dict.get(fencer_d_ref, {"Name": "Unknown", "Nation": "Unknown"})
                        fencer_g = fencer_dict.get(fencer_g_ref, {"Name": "Unknown", "Nation": "Unknown"})

                        if fencer_d["Nation"] == "QAT":
                            review_data[fencer_d["Name"]]["Scored"] += score_d
                            review_data[fencer_d["Name"]]["Conceded"] += score_g
                            review_data[fencer_d["Name"]]["Matches"].append({
                                "Opponent Team": fencer_g["Name"],
                                "Outcome": "Victory" if score_d > score_g else "Defeat" if score_d < score_g else "Draw"
                            })

    with tabs[4]:
        # Final Rankings Tab
        st.header("Final Rankings")
        if not final_rankings_df.empty:
            st.dataframe(final_rankings_df.style.apply(lambda row: ['background-color: yellow' if row['Nation'] == "QAT" else '' for _ in row], axis=1))
        else:
            st.write("No final rankings data available.")

    # Create Word document for download
    buffer = create_word_document(overview_data, qatari_fencers, tables_data, review_data, final_rankings_df)
    st.sidebar.download_button("Download Word Report", data=buffer, file_name="Competition_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
